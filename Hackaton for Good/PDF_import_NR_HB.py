# Hackaton for Peace, Justice and Security
# June 14th - 16th, 2019
# Challenge: National Rapporteur
# Team: NRteam1

# Functions:
# Create list of PDFs (from directory): to get a pathList with all PDFs (as these are provided in subfolders)
# Import PDFs by PyPDF2 from pathList and extract texts to Publications: csv, excel
# Reading Text: from PDF to text (1 file only) by PyPDF2 and extract keywords and look at the weightage (3 methods)
# Import PDFs by PDFMiner from pathList and extract texts to Publications: csv, excel

# Next Actions:
# Extract .doc (or transform .doc to .docx)
# Extract also images, graphics, tables, etc. from the PDFs


# =============================================================================
# Create list of PDFs (from directory)
# =============================================================================
# from: https://stackoverflow.com/questions/3964681/find-all-files-in-a-directory-with-extension-txt-in-python/3964691
import os
## Set up pathList when there are subdirectories and when all one file extension only (e.g. pdf) should be included 
def findFilesInFolder(path, pathList, extension, subFolders = True):
    """  Recursive function to find all files of an extension type in a folder (and optionally in all subfolders too)

    path:        Base directory to find files
    pathList:    A list that stores all paths
    extension:   File extension to find
    subFolders:  Bool.  If True, find files in all subfolders under path. If False, only searches files in the specified folder
    """

    try:   # Trapping a OSError:  File permissions problem I believe
        for entry in os.scandir(path):
            if entry.is_file() and entry.path.endswith(extension):
                pathList.append(entry.path)
            elif entry.is_dir() and subFolders:   # if its a directory, then repeat process as a nested function
                pathList = findFilesInFolder(entry.path, pathList, extension, subFolders)
    except OSError:
        print('Cannot access ' + path +'. Probably a permissions error')

    return pathList

# directory where the PDFs are stored (original data)
#dir_name = r'D:\National Rapporteur Publications' # for PDFs in subfolders 
dir_name = r'D:\FLOW\FLOW Documenten\GitHub\hackathonforgood\downloads_title_url' # for PDFs as scraped from website
extension = ".pdf"

pathList = []
pathList = findFilesInFolder(dir_name, pathList, extension, True) # one document is *.doc
print(pathList)


## Set up pathList when there are no subdirectories and when all file extensions should be included 
import os
def findFilesInFolder(path, pathList):

    try:   # Trapping a OSError:  File permissions problem I believe
        for entry in os.scandir(path):
            if entry.is_file():
                pathList.append(entry.path)
    except OSError:
        print('Cannot access ' + path +'. Probably a permissions error')

    return pathList

# directory where the PDFs are stored (original data)
#dir_name = r'D:\National Rapporteur Publications' # for PDFs in subfolders 
dir_name = r'D:\FLOW\FLOW Documenten\GitHub\hackathonforgood\downloads_title_url' # for PDFs as scraped from website

pathList = []
pathList = findFilesInFolder(dir_name, pathList)
print(pathList)


# =============================================================================
# Import PDFs by PyPDF2 from pathList and extract texts to Publications: csv, excel
# =============================================================================
# based on https://towardsdatascience.com/how-to-extract-keywords-from-pdfs-and-arrange-in-order-of-their-weights-using-python-841556083341
## Import the libraries
import pandas as pd
import numpy as np
import PyPDF2

## Import all PDFs (publications)
publications_pdf = pd.DataFrame()
for i in range(len(pathList)):
    filename = pathList[i]

    pdfFileObj = open(filename,'rb')               # open allows you to read the file
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)   # the pdfReader variable is a readable object that will be parsed
    num_pages = pdfReader.numPages                 # discerning the number of pages will allow us to parse through all the pages

    count = 0
    raw_text = ""

    while count < num_pages:                       
        pageObj = pdfReader.getPage(count)
        count +=1
        raw_text += pageObj.extractText()
    
    publications_pdf.at[i, 'Publication'] = pathList[i]
    publications_pdf.at[i, 'PDF_name'] = pathList[i].rpartition('\\')[2]
    publications_pdf.at[i, 'Raw_Text'] = raw_text 

## Write to file
publications_pdf.to_csv('output/pdf_content.csv', index=False)
publications_pdf.to_excel('output/pdf_content.xlsx', index=False)


# =============================================================================
# Reading Text: from PDF to text (1 file only) by PyPDF2
# =============================================================================
# based on https://towardsdatascience.com/how-to-extract-keywords-from-pdfs-and-arrange-in-order-of-their-weights-using-python-841556083341
import pandas as pd
import numpy as np
import PyPDF2
import re
import nltk
import gensim
import rake_nltk
#import textract # not succeeded to install - but also not needed

# Select filename
#filename = r'D:\National Rapporteur Publications\Statement van de Nationaal Rapporteur ter gelegenheid van de Week tegen Kindermishandeling\Statement van de Nationaal Rapporteur ter gelegenheid van de Week tegen Kindermishandeling .pdf' 
filename = r'D:\National Rapporteur Publications\Vijfde Rapportage Mensenhandel\rapportage-5-(ned)-2006_tcm23-34835.pdf'
#filename = r'D:\National Rapporteur Publications\Achtste Rapportage Mensenhandel\8e rapportage NRM-NL-web_tcm23-34821.pdf' # example file with wrong characters


pdfFileObj = open(filename,'rb')               # open allows you to read the file
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)   # the pdfReader variable is a readable object that will be parsed
num_pages = pdfReader.numPages                 # discerning the number of pages will allow us to parse through all the pages

## Read PDF
count = 0
raw_text = ""
# The while loop will read each page                                                           
while count < num_pages:                       
    pageObj = pdfReader.getPage(count)
    count +=1
    raw_text += pageObj.extractText()

## Text modifications
text = raw_text.replace('\n','')
# Lowercasing each word
text = text.encode('ascii','ignore').lower()
# Change text to string (utf8)
text = text.decode()
print(text)
# remove hyphen '-' between letters, e.g. 'wer-den' to avoid wrong keywords and searching
text = re.sub(r'\b-\b', '', text) # depending on text

## Extracting keywords
keywords = re.findall(r'[a-zA-Z]\w+',text)
len(keywords)                               

# Remove stopwords from keywords
from nltk.corpus import stopwords
dutch_stopwords = set(stopwords.words('dutch'))
keywords = set(keywords) - dutch_stopwords
len(keywords)                               

# List of keywords in DataFrame
df = pd.DataFrame(list(set(keywords)),columns=['keywords']) 

## Calculating weightage
def weightage(word,text,number_of_documents=1):
    word_list = re.findall(word,text)
    number_of_times_word_appeared =len(word_list)
    tf = number_of_times_word_appeared/float(len(text))
    idf = np.log((number_of_documents)/float(number_of_times_word_appeared))
    tf_idf = tf*idf
    return number_of_times_word_appeared, tf, idf, tf_idf

df['number_of_times_word_appeared'] = df['keywords'].apply(lambda x: weightage(x,text)[0])
df['tf'] = df['keywords'].apply(lambda x: weightage(x,text)[1])
df['idf'] = df['keywords'].apply(lambda x: weightage(x,text)[2])
df['tf_idf'] = df['keywords'].apply(lambda x: weightage(x,text)[3])

df = df.sort_values('number_of_times_word_appeared',ascending=False)
df.to_csv('output/keywords.csv')
df.head(25)


## Second Method - Using Gensim library
from gensim.summarization import keywords
import warnings
warnings.filterwarnings("ignore")
values = keywords(text=text,split='\n',scores=True)
data = pd.DataFrame(values,columns=['keyword','score'])
data = data.sort_values('score',ascending=False)
data.head(10)


## Third Approach - Using RAKE (Rapid Automatic Keyword Extraction)
from rake_nltk import Rake
r = Rake()
r.extract_keywords_from_text(text)
phrases = r.get_ranked_phrases_with_scores()
table = pd.DataFrame(phrases,columns=['score','Phrase'])
table = table.sort_values('score',ascending=False)
table.head(10)


# =============================================================================
# Import PDFs by PDFMiner from pathList and extract texts to Publications: csv, excel
# =============================================================================
# based on https://towardsdatascience.com/how-to-extract-keywords-from-pdfs-and-arrange-in-order-of-their-weights-using-python-841556083341
## Import the libraries
import pandas as pd
import numpy as np
from io import StringIO
import pdfminer
import docx

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from PyPDF2 import PdfFileReader

## Convert function PDFMiner 
def convert(fname, pages=None):
    if not pages:
        pagenums = set()
    else:
        pagenums = set(pages)
 
    output = StringIO()
    manager = PDFResourceManager()
    converter = TextConverter(manager, output, laparams=LAParams())
    interpreter = PDFPageInterpreter(manager, converter)
 
    infile = open(fname, 'rb')
    for page in PDFPage.get_pages(infile, pagenums, password="", check_extractable=False):
        interpreter.process_page(page)
    infile.close()
    converter.close()
    text = output.getvalue()
    output.close
    return text

## Convert function getText from doc-file 
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

## Create DataFrame Publications for all PDFs content
publications_pdf = pd.DataFrame(index=range(len(pathList)))
publications_pdf['Publication'] = None
publications_pdf['PDF_name'] = None
publications_pdf['Raw_Text'] = None
publications_pdf['Meta_Data'] = None

# check .doc
for i in range(len(pathList)):
    filename = pathList[i]
    if(filename[-3:] == 'doc'):
        print(i)

for i in range(len(pathList)):
#for i in range(0,10): # for test purposes
    print(i)
    filename = pathList[i]

    publications_pdf.Publication[i] = pathList[i]
    publications_pdf.PDF_name[i] = pathList[i].rpartition('\\')[2]

    # check if pdf extension:
    if(filename[-3:] == 'pdf'):
        pdf = PdfFileReader(open(filename,'rb'))
        nr_pages = pdf.getNumPages()
        pagess = []
        for j in range(nr_pages):
            pagess.append(convert(filename, pages=[j]))
        publications_pdf.Raw_Text[i] = pagess

        # pdf meta data
        publications_pdf.Meta_Data[i] = pdf.getDocumentInfo()

   # check if docx extension:
#    elif(filename[-4:] == 'docx'):
#        document = docx.Document(filename)
#        # extract text
#        publications_pdf.Raw_Text[i] = getText(document)

    # if not a pdf-extension: do not scrape content 
    else:
        print('\nFile with no PDF or DOC extension:', i, '\nFile:', filename)
        publications_pdf.Raw_Text[i] = None


## Write to file
publications_pdf.to_pickle('output/pdfminer190_content_per_page_190619.pkl')
publications_pdf.to_csv('output/pdfminer190_content_per_page_190619.csv', index=False)
publications_pdf.to_excel('output/pdfminer190_content_per_page_190619.xlsx', index=False)

publications_pdf['Publication']

df = publications_pdf.copy()

# =============================================================================
# Text from .doc file (temp & dirty solution)
# =============================================================================
# https://www.nationaalrapporteur.nl/binaries/Persbericht%206de%20rapportage%20Nationaal%20Rapporteur%20Mensenhandel_tcm23-34831.doc

i = 146
filename = 'D:\FLOW\FLOW Documenten\GitHub\hackathonforgood\vanalles\Persbericht 6de rapportage Nationaal Rapporteur Mensenhandel_tcm23-34831.docx' 
publications_pdf.Raw_Text[i] = getText(filename)

# Check metadata: 0https://python-docx.readthedocs.io/en/latest/api/document.html


# =============================================================================
# Read tables from PDF by tabula-py
# =============================================================================
# https://blog.chezo.uno/tabula-py-extract-table-from-pdf-into-python-dataframe-6c7acfa5f302


# =============================================================================
# Import PDF by tika: > does not work yet (API to Java server needed)
# =============================================================================


# =============================================================================
# PDF find on which page a keyword is
# =============================================================================
# https://stackoverflow.com/questions/12571905/finding-on-which-page-a-search-string-is-located-in-a-pdf-document-using-python

#(1) a function to locate the string
def fnPDF_FindText(xFile, xString):
    # xfile : the PDF file in which to look
    # xString : the string to look for
    import pyPdf, re
    PageFound = -1
    pdfDoc = pyPdf.PdfFileReader(file(xFile, "rb"))
    for i in range(0, pdfDoc.getNumPages()):
        content = ""
        content += pdfDoc.getPage(i).extractText() + "\n"
        content1 = content.encode('ascii', 'ignore').lower()
        ResSearch = re.search(xString, content1)
        if ResSearch is not None:
           PageFound = i
           break
    return PageFound
 
#(2) a function to extract the pages of interest    
def fnPDF_ExtractPages(xFileNameOriginal, xFileNameOutput, xPageStart, xPageEnd):
      from pyPdf import PdfFileReader, PdfFileWriter
      output = PdfFileWriter()
      pdfOne = PdfFileReader(file(xFileNameOriginal, "rb"))
      for i in range(xPageStart, xPageEnd):
          output.addPage(pdfOne.getPage(i))
          outputStream = file(xFileNameOutput, "wb")
          output.write(outputStream)
          outputStream.close()


